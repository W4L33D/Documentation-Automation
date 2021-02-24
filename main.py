from docx import Document
from datetime import datetime
from PySimpleGUI import PySimpleGUI as sg

IIT = (
    ('Visual Inspection', 'No cracks, or other damages to the devices or components', 'Same as expected', 'Pass', 'No damages to the module'),
    ('Power on', 'The device should start up and perform the initial instrument boot phase steps', 'All steps were completed within the specified time range', 'Pass', 'Better performance than required'),
    ('IIT Communication', 'RSI values equal 50+-10', 'During the test the value varied from 41-59', 'Pass', 'Two modules had')
)

Temperature = (
    ('Temperature', '40°C ± 0.2°C or 100°F ± 0.4°F', 'noe', 'Pass', 'Within acceptable range'),
    ('', '', '', '', '')
)

IIT_knapp = "IIT"
Temperature_knapp = "Temperature"
Ja = "JA"
Nei = "Nei"
Ferdig = False
tabell = []

while not Ferdig:

    layout = [[sg.Text("Valg av test")], [sg.Button(IIT_knapp)], [sg.Button(Temperature_knapp)]]
    layout2 = [[sg.Text("Ferdig?")], [sg.Button(Ja)], [sg.Button(Nei)]]

    window = sg.Window("Demo", layout, margins=(100, 200))

    while True:
        event, values = window.read()
        if event == IIT_knapp:
            tabell.append(IIT)
            break
        elif event == Temperature_knapp:
            tabell.append(Temperature)
            break
        elif event == sg.WIN_CLOSED:
            break

    window.close()

    window = sg.Window("Demo", layout2, margins=(100, 200))

    while True:
        event, values = window.read()
        if event == Ja:
            Ferdig = True
            break
        elif event == Nei:
            break
        elif event == sg.WIN_CLOSED:
            break
    window.close()

layout = [
    [sg.Text('Please enter your Antall')],
    [sg.Text('Antall', size=(15, 1)), sg.InputText()],
    [sg.Submit(), sg.Cancel()]
]

window = sg.Window('Simple data entry window', layout)
event, values = window.read()
window.close()

print(event, values[0])
lengdetab = len(tabell)
lengde = int(values[0])

now = datetime.now()
dato = now.strftime('%d/%m/%Y')

document = Document()

document.add_heading('Service Rapport', 0)

p = document.add_paragraph('A test for automating document creation.')

document.add_paragraph(
    'Service Engineer: ', style='List Bullet'
)
document.add_paragraph(
    'Date: ' + dato, style='List Bullet'
)
document.add_paragraph(
    'Sted: ', style='List Bullet'
)
document.add_paragraph(
    'Product: ', style='List Bullet'
)


y = 0
for i in range(0, lengde):
    for i in range(0, lengdetab):
        table = document.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Test'
        hdr_cells[1].text = 'Expected'
        hdr_cells[2].text = 'Achieved'
        hdr_cells[3].text = 'Result'
        hdr_cells[4].text = 'Comment'
        for Test, Expected, Achieved, Result, Comment in tabell[y]:
            row_cells = table.add_row().cells
            row_cells[0].text = Test
            row_cells[1].text = Expected
            row_cells[2].text = Achieved
            row_cells[3].text = Result
            row_cells[4].text = Comment
        y = y + 1
    document.add_paragraph(

    )
    y = 0
document.add_page_break()

document.save('Test')